"""Printing module."""

import configparser
import csv
from collections import defaultdict
import datetime
import math
import os
from pathlib import Path
import shutil
import sys
import time

import colorama

import docx
from docx.shared import Mm


# import yagmail

colorama.init()

# epdata_path = Path("d:\\john tillet\\episode_data")
epdata_path = Path("./files")
fees_config = configparser.ConfigParser()
fees_config.read(epdata_path / "FEES.ini")

BILLER = {
    "Dr J Tillett": {
        "name": "Dr John Tillett",
        "address": "7 Henry Lawson Drive, Villawood NSW 2163",
        "provider": "0307195H",
        "abn": "66 781 021 178",
        "contact": "Phone: 8382 6622 Email: john@endoscopy.stvincents.com.au",
        "email": "john.lamia@gmail.com",
    },
    "Dr S Vuong": {
        "name": "Dr Sabine Vuong",
        "address": "PO Box 169 Dulwich Hill 2203",
        "provider": "2349492F",
        "abn": "80 757 898 6622",
        "contact": "8382 3200",
        "email": "sabinevuong@fastmail.com.au",
    },
}


PAGE_NUMBER = 24

today = datetime.datetime.today()
FILE_STR = today.strftime("%y-%m-%d")


def clear():
    print("\033[2J")  # clear screen
    print("\033[1;1H")  # move to top left


def print_account(
    ep, doc, unit, consult_as_float, time_fee, total_fee, biller, page_break=True
):
    """Prints a single accout to a docx document and returns it"""
    biller = BILLER[biller]

    if ep["fund_code"] == "paid":
        doc.add_heading("Tax Receipt for Anaesthetic Fees Paid", level=0)
    else:
        doc.add_heading("Account for Anaesthetic", level=0)

    doc.add_heading(biller["name"], level=2)
    doc.add_heading(biller["address"], level=4)

    if ep["fund_code"] in {"paid", "send_bill"}:
        abn_str = "ABN:  " + (biller["abn"])
        doc.add_heading(
            "Provider Number: " + biller["provider"] + "    " + abn_str, level=5
        )
    else:
        doc.add_heading("Provider Number: " + biller["provider"], level=5)
    doc.add_heading(biller["contact"], level=5)
    doc.add_paragraph("")

    h_head = doc.add_heading("Patient Details", level=4)
    inv_string = "%s%s%s" % (("  " * 20), "Invoice Number  ", ep["invoice"])
    h_head.add_run(inv_string)
    doc.add_paragraph("")
    a1 = doc.add_paragraph()
    a1.add_run("%s" % ep["name"]).bold = True
    doc.add_paragraph("%s" % ep["address"])
    doc.add_paragraph("Date of birth  %s" % ep["dob"])
    if ep["fund_code"] in {"ga", "adf"}:
        doc.add_paragraph(
            "Fund:  %s   Episode Number:  %s" % (ep["fund_name"], ep["ref"])
        )
        p_ga = doc.add_paragraph("Approval Number:")
        ga_str = "   %s" % ep["fund_number"]
        p_ga.add_run(ga_str)

    elif ep["fund_code"] == "paid":
        pass
    else:
        doc.add_paragraph(
            "Fund:  %s   Number:  %s" % (ep["fund_name"], ep["fund_number"])
        )
        if ep["fund_code"] == "send_bill":
            doc.add_paragraph("Patient not registered with Medicare for this service.")
        elif ep["fund_code"] == "paid":
            pass
        else:
            p_mc = doc.add_paragraph("Medicare Number")
            mc_str = "   %s    ref  %s" % (ep["medicare_no"], ep["ref"])
            p_mc.add_run(mc_str)
    doc.add_paragraph("Date of Procedure:  %s" % ep["date"])
    doc.add_paragraph("Procedure performed by %s" % ep["doctor"])
    doc.add_paragraph("Diagnostic Endoscopy Centre, Darlinghurst, NSW 2010")
    doc.add_paragraph("Facility ID 657131A")

    doc.add_paragraph("Item Number%sFee" % (" " * 10))

    p_cons = doc.add_paragraph("17610")
    cons_str = "%.2f" % consult_as_float
    cons_str = cons_str.rjust(25)
    p_cons.add_run(cons_str)

    if ep["upper"] in {"Yes", "upper_Yes"}:
        p_endo = doc.add_paragraph("20740")
        endo_str = "%.2f" % (unit * 5)
        endo_str = endo_str.rjust(24)
        p_endo.add_run(endo_str)
        if ep["lower"] in {"Yes", "lower_Yes"}:
            p_col = doc.add_paragraph("20810")
            col_str = "%.2f" % 0.0
            col_str = col_str.rjust(26)
            p_col.add_run(col_str)
    if ep["upper"] in {"No", "upper_No"} and ep["lower"] in {"Yes", "lower_Yes"}:
        p_col = doc.add_paragraph("20810")
        col_str = "%.2f" % (unit * 4)
        col_str = col_str.rjust(24)
        p_col.add_run(col_str)
    if ep["seventy"] in {"Yes", "age70_Yes"}:
        p_age = doc.add_paragraph("25014")
        age_str = "%.2f" % unit
        age_str = age_str.rjust(25)
        p_age.add_run(age_str)
    if ep["asa_3"] in {"Yes", "asa3_Yes"}:
        p_sick = doc.add_paragraph("25000")
        sick_str = "%.2f" % unit
        sick_str = sick_str.rjust(25)
        p_sick.add_run(sick_str)
    if ep["asa_3"] in {"asa3_Four"}:
        p_sick = doc.add_paragraph("25005")
        sick_str = "%.2f" % unit
        sick_str = sick_str.rjust(25)
        p_sick.add_run(sick_str)

    p_time_fee = doc.add_paragraph(ep["time"])
    time_fee_str = "%.2f" % time_fee
    time_fee_str = time_fee_str.rjust(25)
    p_time_fee.add_run(time_fee_str)

    p_tot = doc.add_paragraph("Total Fee")
    tot_str = "$%.2f" % total_fee
    tot_str = tot_str.rjust(19)
    p_tot.add_run(tot_str).bold = True

    if ep["fund_code"] == "paid":
        p_tot = doc.add_paragraph("Paid")
        tot_str = "$%.2f" % total_fee
        tot_str = tot_str.rjust(26)
        p_tot.add_run(tot_str).bold = True

    if ep["fund_code"] == "paid":
        p_tot = doc.add_paragraph("Owing")
        tot_str = "$0.00"
        tot_str = tot_str.rjust(25)
        p_tot.add_run(tot_str).bold = True

    doc.add_paragraph("")
    p_gst = doc.add_paragraph("")
    p_gst.add_run("No item on this invoice attracts GST").italic = True
    if page_break:
        doc.add_page_break()

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    return doc


def print_calc(n):
    """calculate number of invoices to print if there are more than 20."""
    divisor = math.ceil(n / 20)
    return math.floor(n / divisor)


# FUND_CODES = [
#     "Pay Later",
#     "paid",
#     "send_bill",
#     "va",
#     "bb",
#     "hcf",
#     "bup",
#     "mpl",
#     "nib",
#     "ahm",
#     "ga",
#     "gu",
#     "ama",
#     "ahsa",
# ]


def format_fixed_width(rows):
    """Taken from Trey Hunner's Python Morsels. Formats output for summary."""
    column_lengths = [max(len(cell) for cell in col) for col in zip(*rows)]
    output = ""
    for row in rows:
        for column, length in zip(row, column_lengths):
            output += column.ljust(length + 2)
        output = output.rstrip() + "\n"
    return output.rstrip()


def mail_and_backup(anaesthetist, file_type):
    surname = anaesthetist.split()[-1]
    yag = yagmail.SMTP("john.lamia@gmail.com", "amjteytpugcrmuhw")
    to = BILLER[anaesthetist]["email"]

    body = "body test"
    if file_type == "summary":
        path = "d:/john tillet/episode_data/sedation/accts_summary.txt"
        subject = "DEC billing summary"
        body = "Attached is a summary of your acounts proccesed today"
        save_destination = "d:/john tillet/episode_data/sedation/backup/{}-{}-accts_summary.txt".format(
            FILE_STR, surname
        )
    elif file_type == "csv":
        path = "d:/john tillet/episode_data/sedation/{}.csv".format(surname)
        subject = "DEC billing csv"
        body = "Attached is the csv file of your acounts proccesed today"
        save_destination = (
            "d:/john tillet/episode_data/sedation/backup/{}-{}-csv.csv".format(
                FILE_STR, surname
            )
        )
    elif file_type == "docx":
        path = "d:/john tillet/episode_data/sedation/accts.docx"
        subject = "DEC accounts"
        body = "Attached are your acounts proccesed today"
        save_destination = (
            "d:/john tillet/episode_data/sedation/backup/{}-{}-accts.docx".format(
                FILE_STR, surname
            )
        )

    path = os.path.realpath(path)
    save_path = os.path.realpath(save_destination)
    content = [path]
    yag.send(to, subject, body, content)
    shutil.copy(path, save_path)


def sort_funds(x):
    if x == "send_bill":
        return 10
    elif x == "adf":
        return 15
    elif x == "paid":
        return 20
    elif x == "bb":
        return 30
    elif x == "va":
        return 40
    elif x == "hcf":
        return 50
    elif x == "bup":
        return 60
    elif x == "mpl":
        return 70
    elif x == "nib":
        return 80
    elif x == "ahm":
        return 90
    elif x == "ga":
        return 100
    elif x == "gu":
        return 175
    elif x == "ama":
        return 110
    elif x in {"lt", "hh", "sl"}:
        return 165
    else:
        return 150


def process_acc(grand_total, ep):
    """Do some calculations before printing account.

    Input - grand_total a float and ep an Ordered Dictionary
    """

    # get fees from module depending on fund
    #    fee_package = FUND_FEES[ep["fund_code"]]
    #    consult_as_float = float(fee_package[0])
    #    unit = float(fee_package[1])
    consult_as_float = float(fees_config[ep["fund_code"]]["consult"])
    unit = float(fees_config[ep["fund_code"]]["unit"])

    # get time info and calculate time fee
    # the fourth digit in the time code gives the number of units
    time_length = int(ep["time"][3])
    time_fee = time_length * unit

    # calculate total_fee, initialise total_fee with consult
    total_fee = consult_as_float

    if ep["upper"] in {"Yes", "upper_Yes"}:
        total_fee += unit * 5
    if ep["upper"] in {"No", "upper_No"} and ep["lower"] in {"Yes", "lower_Yes"}:
        total_fee += unit * 4
    if ep["seventy"] in {"Yes", "age70_Yes"}:
        total_fee += unit
    if ep["asa_3"] in {"Yes", "asa3_Yes"}:
        total_fee += unit
    if ep["asa_3"] in {"asa3_Four"}:
        total_fee += unit * 2

    # add on time fees
    total_fee = total_fee + (time_length * unit)
    grand_total += total_fee

    return (grand_total, consult_as_float, unit, time_fee, total_fee)


def cleanup(datafile, masterfile, summaryfile, printfile, anaesthetist):
    """Close and delete files at end of script."""
    input("Press Enter to open accounts summary.")
    os.startfile(summaryfile)
    print()
    input("Press Enter to open accounts file.")
    os.startfile(printfile)
    time.sleep(3)

    print()
    while True:
        merge = """All done!

Press y to merge csv with masterfile.

Press n to just exit (current csv file will stay in place.)

"""
        flag = input(merge)
        if flag in {"y", "n"}:
            break
    if flag == "n":
        pass

    elif flag == "y":
        with open(datafile) as csvhandle:
            with open(masterfile, "a") as filehandle:
                csvdata = csv.reader(csvhandle)
                csvwriter = csv.writer(filehandle, dialect="excel", lineterminator="\n")
                for p in csvdata:
                    csvwriter.writerow(p)
        try:
            os.remove(datafile)
        except FileNotFoundError:
            pass


# Print iterations progress
def printProgressBar(
    iteration, total, prefix="", suffix="", decimals=1, length=25, fill="█"
):
    """
    Call in a loop to create terminal progress bar
    @params:
        iteration   - Required  : current iteration (Int)
        total       - Required  : total iterations (Int)
        prefix      - Optional  : prefix string (Str)
        suffix      - Optional  : suffix string (Str)
        decimals    - Optional  : positive number of decimals in percent complete (Int)
        length      - Optional  : character length of bar (Int)
        fill        - Optional  : bar fill character (Str)
    """
    percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
    filledLength = int(length * iteration // total)
    bar = fill * filledLength + "-" * (length - filledLength)
    print("\r%s |%s| %s%% %s" % (prefix, bar, percent, suffix), end="\r")
    # Print New Line on Complete
    if iteration == total:
        print()


def main():
    """Print accounts in batches of funds."""

    biller = "Dr J Tillett"

    clear()
    print_set = {"adf", "send_bill"}
    biller_surmame = biller.split()[-1]
    base = "D:\\JOHN TILLET\\episode_data\\sedation\\"
    summaryfile = base + "accts_summary.txt"
    printfile = base + "accts.docx"
    datafile = base + biller_surmame + ".csv"
    masterfile = base + biller_surmame + "_master.csv"

    headers = [
        "date",
        "name",
        "address",
        "dob",
        "medicare_no",
        "ref",
        "fund_name",
        "fund_number",
        "fund_code",
        "doctor",
        "upper",
        "lower",
        "seventy",
        "asa_3",
        "time",
        "invoice",
    ]
    try:
        os.remove(summaryfile)
    except IOError:
        pass

    try:
        with open(datafile) as csvfile:
            reader = csv.DictReader(csvfile, headers)
            ep_list = list(reader)

    except IOError:
        print("No csv file found.")
        sys.exit(1)

    clear()
    print("Printing Dr {}'s accounts".format(biller_surmame))
    print()

    doc = docx.Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Verdana"

    # build a dictionary pat_dict where keys are fund names
    # and values are a list of episodes
    pat_dict = defaultdict(list)

    for episode in ep_list:
        if episode["fund_code"] in print_set:
            fund_id = episode["fund_code"]
            pat_dict[fund_id].append(episode)

    summary_list = []
    # for each fund in pat_dict get the list of episodes ep_list and print
    # them out in equal batches < 20
    length = 0
    for fund in pat_dict:
        length += len(pat_dict[fund])

    iteration = 0
    printProgressBar(
        iteration, length, prefix="Progress:", suffix="Complete", length=20
    )

    for fu in sorted(pat_dict.keys(), reverse=False, key=sort_funds):
        ep_list = pat_dict[fu]

        fund_len = len(ep_list)
        fund_left = fund_len
        start = 0
        if fund_len == 0:
            continue
        while True:
            grand_total = 0.0
            num_p = print_calc(fund_left)
            end = start + num_p
            for np in ep_list[start:end]:
                iteration += 1
                processed = process_acc(grand_total, np)
                grand_total, consult_as_float, unit, time_fee, total_fee = processed
                acc = print_account(
                    np, doc, unit, consult_as_float, time_fee, total_fee, biller
                )

            #            with open(summaryfile, 'a') as file:
            #                if np['fund_code'] == 'os':
            #                    file.write('Overseas -- %s\n ' % num_p)
            #                else:
            #                    file.write('%s -- %s\n ' % (fu, num_p))

            if np["fund_code"] == "send_bill":
                summary_name = "Not paid yet"
            elif np["fund_code"] == "paid":
                summary_name = "Paid at DEC"
            else:
                summary_name = np["fund_name"]

            summary_list.append(list([summary_name, str(num_p)]))

            # acc = print_batch_header(doc, fu, num_p, grand_total, headers_datafile)
            fund_left -= num_p
            start += num_p

            printProgressBar(
                iteration, length, prefix="Progress:", suffix="Complete", length=20
            )
            time.sleep(0.1)
            if start >= fund_len:
                break
    with open(summaryfile, "wt") as file:
        file.write(format_fixed_width(summary_list))
    acc.save(printfile)
    print()

    mail_and_backup(biller, "summary")
    print()
    print("Emailing the accounts..")
    mail_and_backup(biller, "docx")
    print()

    cleanup(datafile, masterfile, summaryfile, printfile, biller)


if __name__ == "__main__":
    main()
